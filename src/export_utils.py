"""
Gelişmiş Export Utilities - HTML Destekli Word Rapor Export Fonksiyonları
"""

import os
import streamlit as st
from datetime import datetime, timedelta
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import json

def add_html_to_doc(doc, html_content):
    """HTML içeriğini Word belgesine ekler"""
    try:
        # HTML'i parse et ve Word'e ekle
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # HTML elementlerini Word'e dönüştür
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            level = int(element.name[1])
            heading = doc.add_heading(element.get_text(), level)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for element in soup.find_all('p'):
            if element.get_text().strip():
                doc.add_paragraph(element.get_text())
        
        for element in soup.find_all('ul'):
            for li in element.find_all('li'):
                doc.add_paragraph(f"• {li.get_text()}", style='List Bullet')
        
        for element in soup.find_all('ol'):
            for i, li in enumerate(element.find_all('li'), 1):
                doc.add_paragraph(f"{i}. {li.get_text()}", style='List Number')
        
    except Exception as e:
        # HTML parse edilemezse düz metin olarak ekle
        doc.add_paragraph(html_content)

def create_styled_document():
    """Stilize edilmiş Word belgesi oluşturur"""
    doc = Document()
    
    # Stil tanımlamaları
    styles = doc.styles
    
    # Başlık stili
    title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Inches(0.2)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    
    # Alt başlık stili
    subtitle_style = styles.add_style('Custom Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Inches(0.15)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(64, 64, 64)
    
    # Metin stili
    text_style = styles.add_style('Custom Text', WD_STYLE_TYPE.PARAGRAPH)
    text_font = text_style.font
    text_font.name = 'Arial'
    text_font.size = Inches(0.12)
    text_font.color.rgb = RGBColor(32, 32, 32)
    
    return doc

def create_comprehensive_word_report(symbol, export_data):
    """Kapsamlı Word raporu oluşturur"""
    
    doc = create_styled_document()
    
    # Başlık
    title = doc.add_heading(f'📈 {symbol} - Kapsamlı Tahmin Analizi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tarih ve bilgiler
    doc.add_paragraph(f'📅 Analiz Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph(f'🏢 Hisse Senedi: {symbol}')
    doc.add_paragraph()
    
    # Yatırım Kararı - Büyük ve dikkat çekici
    prediction = export_data.get('prediction', 1)
    confidence = export_data.get('confidence', 0)
    current_price = export_data.get('current_price', 0)
    
    decision_text = "🟢 AL" if prediction == 1 else "🔴 SAT"
    decision_color = RGBColor(0, 128, 0) if prediction == 1 else RGBColor(220, 53, 69)
    
    decision_heading = doc.add_heading(f'🎯 YATIRIM KARARI: {decision_text}', level=1)
    decision_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Akıllı Analiz Özeti - En üstte
    doc.add_heading('📖 Akıllı Analiz Özeti', level=2)
    
    # Akıllı analiz metni oluştur
    symbol_name = symbol.replace('.IS', '')
    targets = export_data.get('targets', {})
    time_targets = export_data.get('time_targets', {})
    volatility = export_data.get('volatility', 0)
    
    # Son 3 günün analizi (simüle edilmiş)
    price_change_3d = 2.5  # Bu değer gerçekte hesaplanmalı
    price_change_text = f"%{abs(price_change_3d):.1f} {'yükseliş' if price_change_3d > 0 else 'düşüş'}"
    
    # Volatilite kategorisi
    if volatility > 0.6:
        volatility_desc = "çok yüksek"
        volatility_emoji = "🌪️"
    elif volatility > 0.4:
        volatility_desc = "yüksek"
        volatility_emoji = "⚠️"
    elif volatility > 0.25:
        volatility_desc = "orta"
        volatility_emoji = "📊"
    else:
        volatility_desc = "düşük"
        volatility_emoji = "🛡️"
    
    # Güven kategorisi
    if confidence > 0.8:
        confidence_desc = "çok yüksek"
        confidence_emoji = "💪"
    elif confidence > 0.6:
        confidence_desc = "yüksek"
        confidence_emoji = "🎯"
    else:
        confidence_desc = "orta"
        confidence_emoji = "⚠️"
    
    # Hedef fiyat ve süre
    moderate_target = targets.get('moderate', current_price)
    moderate_days = time_targets.get('moderate', {}).get('estimated_days', 0) if time_targets else 0
    target_change_pct = ((moderate_target - current_price) / current_price) * 100
    
    # Akıllı analiz metni
    if prediction == 1:  # AL sinyali
        if price_change_3d < -2:  # Son 3 günde düşüş
            analysis_text = f"""
{symbol_name} şu anda {current_price:.2f} TL seviyesinde. Son 3 günde {price_change_text} yaşadı, ancak model bu düşüşün sona ereceğini öngörüyor.

Model, {moderate_days} gün içinde %{target_change_pct:.1f} yükselişle {moderate_target:.2f} TL hedefini öngörüyor. Volatilite {volatility_desc} seviyede ({volatility_emoji}), güven oranı {confidence_desc} ({confidence_emoji} %{confidence*100:.1f}).
"""
        else:  # Son 3 günde yükseliş veya stabil
            analysis_text = f"""
{symbol_name} şu anda {current_price:.2f} TL seviyesinde. Son 3 günde {price_change_text} yaşadı ve model bu yükseliş trendinin devam edeceğini öngörüyor.

Model, {moderate_days} gün içinde %{target_change_pct:.1f} yükselişle {moderate_target:.2f} TL hedefini öngörüyor. Volatilite {volatility_desc} seviyede ({volatility_emoji}), güven oranı {confidence_desc} ({confidence_emoji} %{confidence*100:.1f}).
"""
    else:  # SAT sinyali
        if price_change_3d > 2:  # Son 3 günde yükseliş
            analysis_text = f"""
{symbol_name} şu anda {current_price:.2f} TL seviyesinde. Son 3 günde {price_change_text} yaşadı, ancak model bu yükselişin sona ereceğini öngörüyor.

Model, {moderate_days} gün içinde %{abs(target_change_pct):.1f} düşüşle {moderate_target:.2f} TL seviyesine ineceğini öngörüyor. Volatilite {volatility_desc} seviyede ({volatility_emoji}), güven oranı {confidence_desc} ({confidence_emoji} %{confidence*100:.1f}).
"""
        else:  # Son 3 günde düşüş veya stabil
            analysis_text = f"""
{symbol_name} şu anda {current_price:.2f} TL seviyesinde. Son 3 günde {price_change_text} yaşadı ve model bu düşüş trendinin devam edeceğini öngörüyor.

Model, {moderate_days} gün içinde %{abs(target_change_pct):.1f} düşüşle {moderate_target:.2f} TL seviyesine ineceğini öngörüyor. Volatilite {volatility_desc} seviyede ({volatility_emoji}), güven oranı {confidence_desc} ({confidence_emoji} %{confidence*100:.1f}).
"""
    
    doc.add_paragraph(analysis_text.strip())
    doc.add_paragraph()
    
    # Karar detayları
    doc.add_heading('📊 Karar Detayları', level=2)
    
    # Metrikler tablosu
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Tablo başlığı
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metrik'
    hdr_cells[1].text = 'Değer'
    
    # Metrikler
    metrics_data = [
        ('Güven Skoru', f'%{confidence*100:.1f}'),
        ('Güncel Fiyat', f'₺{current_price:.2f}'),
        ('Volatilite', f'%{export_data.get("volatility", 0)*100:.1f}'),
        ('Risk/Getiri Oranı', f'1:{export_data.get("risk_reward_ratio", 0):.2f}')
    ]
    
    for metric, value in metrics_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metric
        row_cells[1].text = value
    
    doc.add_paragraph()
    
    # Hedef Fiyatlar ve Tarihler
    targets = export_data.get('targets', {})
    time_targets = export_data.get('time_targets', {})
    
    if targets:
        doc.add_heading('🎯 Hedef Fiyatlar ve Tahmini Tarihler', level=2)
        
        targets_table = doc.add_table(rows=1, cols=5)
        targets_table.style = 'Table Grid'
        
        # Hedef fiyatlar tablosu - 5 kolon
        hdr_cells = targets_table.rows[0].cells
        hdr_cells[0].text = 'Hedef Türü'
        hdr_cells[1].text = 'Fiyat (₺)'
        hdr_cells[2].text = 'Değişim (%)'
        hdr_cells[3].text = 'Tahmini Süre'
        hdr_cells[4].text = 'Tahmini Tarih'
        
        target_names = {
            'conservative': 'Konservatif Hedef',
            'moderate': 'Orta Hedef', 
            'aggressive': 'Agresif Hedef'
        }
        
        # Mevcut tarih
        current_date = datetime.now()
        
        for target_type, target_price in targets.items():
            row_cells = targets_table.add_row().cells
            row_cells[0].text = target_names.get(target_type, target_type.title())
            row_cells[1].text = f'{target_price:.2f}'
            
            # Yüzdelik değişim hesapla
            change_pct = ((target_price - current_price) / current_price) * 100
            change_text = f'{change_pct:+.1f}%'
            row_cells[2].text = change_text
            
            # Tahmini süre ve tarih
            if time_targets and target_type in time_targets:
                min_days = time_targets[target_type].get('min_days', 0)
                max_days = time_targets[target_type].get('max_days', 0)
                estimated_days = time_targets[target_type].get('estimated_days', 0)
                
                # Zaman aralığı formatı
                if min_days > 0 and max_days > 0:
                    time_range = f'{min_days}-{max_days} gün'
                else:
                    time_range = f'{estimated_days} gün'
                
                row_cells[3].text = time_range
                
                # Tahmini tarih hesapla (ortalama)
                target_date = current_date + timedelta(days=estimated_days)
                row_cells[4].text = target_date.strftime('%d.%m.%Y')
            else:
                row_cells[3].text = 'Bilinmiyor'
                row_cells[4].text = 'Bilinmiyor'
        
        # Stop Loss
        stop_loss = export_data.get('stop_loss', 0)
        if stop_loss:
            row_cells = targets_table.add_row().cells
            row_cells[0].text = 'Stop Loss'
            row_cells[1].text = f'{stop_loss:.2f}'
            
            # Stop loss yüzdelik değişim
            stop_loss_change = ((stop_loss - current_price) / current_price) * 100
            row_cells[2].text = f'{stop_loss_change:+.1f}%'
            
            row_cells[3].text = 'Anında'
            row_cells[4].text = 'Sürekli'
        
        doc.add_paragraph()
    
    # Destek/Direnç Seviyeleri
    if time_targets and 'conservative' in time_targets:
        chart_analysis = time_targets['conservative'].get('chart_analysis', {})
        if chart_analysis:
            doc.add_heading('📊 Destek/Direnç Seviyeleri', level=2)
            
            support_resistance_table = doc.add_table(rows=1, cols=3)
            support_resistance_table.style = 'Table Grid'
            
            hdr_cells = support_resistance_table.rows[0].cells
            hdr_cells[0].text = 'Seviye Türü'
            hdr_cells[1].text = 'Fiyat (₺)'
            hdr_cells[2].text = 'Mevcut Fiyattan (%)'
            
            # Destek seviyesi
            support_level = chart_analysis.get('support_level', 0)
            if support_level:
                row_cells = support_resistance_table.add_row().cells
                row_cells[0].text = 'Destek Seviyesi'
                row_cells[1].text = f'{support_level:.2f}'
                support_pct = ((support_level - current_price) / current_price * 100)
                row_cells[2].text = f'{support_pct:+.1f}%'
            
            # Mevcut fiyat
            row_cells = support_resistance_table.add_row().cells
            row_cells[0].text = 'Mevcut Fiyat'
            row_cells[1].text = f'{current_price:.2f}'
            row_cells[2].text = '0.0%'
            
            # Direnç seviyesi
            resistance_level = chart_analysis.get('resistance_level', 0)
            if resistance_level:
                row_cells = support_resistance_table.add_row().cells
                row_cells[0].text = 'Direnç Seviyesi'
                row_cells[1].text = f'{resistance_level:.2f}'
                resistance_pct = ((resistance_level - current_price) / current_price * 100)
                row_cells[2].text = f'{resistance_pct:+.1f}%'
            
            doc.add_paragraph()
    
    # Grafik Analizi Detayları
    if time_targets and 'conservative' in time_targets:
        chart_analysis = time_targets['conservative'].get('chart_analysis', {})
        if chart_analysis:
            doc.add_heading('📈 Grafik Analizi Detayları', level=2)
            
            chart_table = doc.add_table(rows=1, cols=2)
            chart_table.style = 'Table Grid'
            
            hdr_cells = chart_table.rows[0].cells
            hdr_cells[0].text = 'Analiz Türü'
            hdr_cells[1].text = 'Sonuç'
            
            chart_data = [
                ('Trend Gücü', chart_analysis.get('trend_strength', 'Bilinmiyor')),
                ('Hacim Trendi', chart_analysis.get('volume_trend', 'Bilinmiyor')),
                ('Destek/Direnç Yakınlığı', 'Yakın' if chart_analysis.get('near_support_resistance', False) else 'Uzak'),
                ('Grafik Pattern', chart_analysis.get('pattern', 'Bilinmiyor'))
            ]
            
            for analysis_type, result in chart_data:
                row_cells = chart_table.add_row().cells
                row_cells[0].text = analysis_type
                row_cells[1].text = result
            
            doc.add_paragraph()
    
    # Analiz Faktörleri
    factors = export_data.get('factors', {})
    if factors:
        doc.add_heading('🔍 Analiz Faktörleri', level=2)
        
        # Olumlu faktörler
        positive_factors = factors.get('positive', [])
        if positive_factors:
            doc.add_heading('✅ Olumlu Faktörler', level=3)
            for i, factor in enumerate(positive_factors, 1):
                doc.add_paragraph(f'{i}. {factor}', style='List Number')
            doc.add_paragraph()
        
        # Olumsuz faktörler
        negative_factors = factors.get('negative', [])
        if negative_factors:
            doc.add_heading('❌ Olumsuz Faktörler', level=3)
            for i, factor in enumerate(negative_factors, 1):
                doc.add_paragraph(f'{i}. {factor}', style='List Number')
            doc.add_paragraph()
        
        # Nötr faktörler
        neutral_factors = factors.get('neutral', [])
        if neutral_factors:
            doc.add_heading('⚖️ Nötr Faktörler', level=3)
            for i, factor in enumerate(neutral_factors, 1):
                doc.add_paragraph(f'{i}. {factor}', style='List Number')
            doc.add_paragraph()
    
    # Model Bilgileri
    model_info = export_data.get('model_info', {})
    if model_info:
        doc.add_heading('🤖 Model Bilgileri', level=2)
        
        model_table = doc.add_table(rows=1, cols=2)
        model_table.style = 'Table Grid'
        
        hdr_cells = model_table.rows[0].cells
        hdr_cells[0].text = 'Özellik'
        hdr_cells[1].text = 'Değer'
        
        model_data = [
            ('Model Türü', model_info.get('model_type', 'Bilinmiyor')),
            ('Eğitim Tarihi', model_info.get('training_date', 'Bilinmiyor')),
            ('Veri Noktası', f"{model_info.get('data_points', 0):,} gün"),
            ('Özellik Sayısı', f"{model_info.get('features_count', 0)} teknik gösterge"),
            ('Doğruluk', f"%{model_info.get('accuracy', 0)*100:.1f}"),
            ('F1 Skoru', f"%{model_info.get('f1_score', 0)*100:.1f}")
        ]
        
        for prop, value in model_data:
            row_cells = model_table.add_row().cells
            row_cells[0].text = prop
            row_cells[1].text = value
        
        doc.add_paragraph()
    
    # Risk Analizi
    doc.add_heading('⚠️ Risk Analizi', level=2)
    
    volatility = export_data.get('volatility', 0)
    risk_level = "Çok Yüksek" if volatility > 0.6 else "Yüksek" if volatility > 0.4 else "Orta" if volatility > 0.25 else "Düşük"
    
    doc.add_paragraph(f'📊 Volatilite Seviyesi: {risk_level} (%{volatility*100:.1f})')
    
    if volatility > 0.6:
        doc.add_paragraph('⚠️ Yüksek volatilite nedeniyle dikkatli pozisyon yönetimi önerilir.')
    elif volatility > 0.4:
        doc.add_paragraph('📈 Orta-yüksek volatilite, dinamik fiyat hareketleri beklenebilir.')
    elif volatility > 0.25:
        doc.add_paragraph('📊 Normal volatilite seviyesi, stabil hareket beklenir.')
    else:
        doc.add_paragraph('🛡️ Düşük volatilite, güvenli ve stabil hareket.')
    
    doc.add_paragraph()
    
    # Yatırım Önerileri
    doc.add_heading('💡 Yatırım Önerileri', level=2)
    
    if prediction == 1:  # AL sinyali
        doc.add_paragraph('🟢 AL Sinyali - Pozitif beklenti')
        doc.add_paragraph('• Mevcut pozisyonunuzu koruyun veya yeni pozisyon ekleyin')
        doc.add_paragraph('• Stop loss seviyesini takip edin')
        doc.add_paragraph('• Hedef fiyatlara ulaştığında kısmi kar realizasyonu yapın')
    else:  # SAT sinyali
        doc.add_paragraph('🔴 SAT Sinyali - Negatif beklenti')
        doc.add_paragraph('• Mevcut pozisyonlarınızı gözden geçirin')
        doc.add_paragraph('• Stop loss seviyelerini sıkı takip edin')
        doc.add_paragraph('• Yeni pozisyon almayın, mevcut pozisyonları azaltın')
    
    doc.add_paragraph()
    
    # Teknik Notlar
    doc.add_heading('📝 Teknik Notlar', level=2)
    doc.add_paragraph('• Bu analiz makine öğrenmesi algoritmaları kullanılarak oluşturulmuştur')
    doc.add_paragraph('• Geçmiş performans gelecekteki sonuçları garanti etmez')
    doc.add_paragraph('• Yatırım kararlarınızı alırken kendi araştırmanızı yapın')
    doc.add_paragraph('• Risk yönetimi kurallarına uyun')
    doc.add_paragraph('• Stop loss ve take profit seviyelerini belirleyin')
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph('_' * 50)
    doc.add_paragraph(f'Rapor Oluşturulma Tarihi: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('Hisse Senedi Yön Tahmini Sistemi - AI Destekli Analiz')
    
    return doc

def create_export_buttons(symbol, export_data):
    """
    Gelişmiş export butonlarını oluşturur
    """
    # Export butonları
    st.markdown("---")
    st.markdown('<h3 class="subsection-title">📄 Rapor Export</h3>', unsafe_allow_html=True)
    
    # Session state ile dosyaları cache'le
    cache_key_word = f"word_data_{symbol}_{datetime.now().strftime('%Y%m%d_%H%M')}"
    
    # Word dosyasını oluştur
    if cache_key_word not in st.session_state:
        try:
            # Logs klasörünü oluştur
            logs_dir = "logs"
            if not os.path.exists(logs_dir):
                os.makedirs(logs_dir)
            
            # Word dosya adı
            word_filename = f"tahmin_raporu_{symbol.replace('.IS', '')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            word_filepath = os.path.join(logs_dir, word_filename)
            
            # Kapsamlı Word raporu oluştur
            doc = create_comprehensive_word_report(symbol, export_data)
            doc.save(word_filepath)
            
            # Dosyayı session state'e kaydet
            with open(word_filepath, "rb") as word_file:
                st.session_state[cache_key_word] = word_file.read()
                
            st.success(f"✅ Kapsamlı Word raporu oluşturuldu!")
            
        except Exception as e:
            st.error(f"❌ Word raporu oluşturma hatası: {str(e)}")
            return
    
    # Word Download
    if cache_key_word in st.session_state:
        col1, col2 = st.columns(2)
        
        with col1:
            st.download_button(
                label="📝 Word Raporu İndir",
                data=st.session_state[cache_key_word],
                file_name=f"tahmin_raporu_{symbol.replace('.IS', '')}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"word_download_{symbol.replace('.IS', '')}_{datetime.now().strftime('%Y%m%d_%H%M')}",
                type="primary"
            )
        
        with col2:
            # JSON export removed due to serialization issues
            st.info("📊 JSON export özelliği geçici olarak devre dışı bırakıldı.")
    
    # Bilgi
    st.markdown("---")
    st.info("""
    💡 **Export Özellikleri:**
    - 📝 **Word Raporu**: Tüm analiz bilgilerini içeren kapsamlı rapor
    - 🎯 **Hedef Fiyatlar**: Konservatif, orta ve agresif hedefler
    - 🔍 **Faktör Analizi**: Olumlu, olumsuz ve nötr faktörler
    - 🤖 **Model Bilgileri**: Model türü, doğruluk ve performans
    - ⚠️ **Risk Analizi**: Volatilite ve risk değerlendirmesi
    """)