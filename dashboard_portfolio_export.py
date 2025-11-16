"""
Portföy Yöneticisi Export Fonksiyonları
"""

import os
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE


def create_portfolio_recommendations_export(recommendations, portfolio, export_date=None):
    """Portföy önerilerini Word formatında export eder"""
    
    if export_date is None:
        export_date = datetime.now()
    
    doc = Document()
    
    # Başlık
    title = doc.add_heading('🤖 Robot Portföy Yöneticisi - Günlük Öneriler Raporu', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tarih
    doc.add_paragraph(f'📅 Rapor Tarihi: {export_date.strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph()
    
    # Portföy Özeti
    doc.add_heading('📊 Portföy Özeti', level=1)
    
    total_cash = portfolio.get('cash', 0)
    portfolio_value = sum([s['quantity'] * s['avg_cost'] for s in portfolio.get('stocks', {}).values()])
    total_capital = total_cash + portfolio_value
    
    portfolio_table = doc.add_table(rows=4, cols=2)
    portfolio_table.style = 'Table Grid'
    
    portfolio_data = [
        ('💵 Toplam Nakit', f'{total_cash:,.2f} TL'),
        ('📊 Portföy Değeri', f'{portfolio_value:,.2f} TL'),
        ('💰 Toplam Sermaye', f'{total_capital:,.2f} TL'),
        ('📈 Hisse Sayısı', f'{len(portfolio.get("stocks", {}))} adet')
    ]
    
    for i, (label, value) in enumerate(portfolio_data):
        portfolio_table.rows[i].cells[0].text = label
        portfolio_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # İşlem önerilerini grupla
    sell_actions = [r for r in recommendations if r['action'] in ['SAT', 'KISMEN SAT']]
    buy_actions = [r for r in recommendations if r['action'] == 'YENİ AL']
    increase_actions = [r for r in recommendations if r['action'] == 'ARTIR']
    hold_actions = [r for r in recommendations if r['action'] == 'TUT']
    
    # 1. SATIŞ ÖNERİLERİ
    if sell_actions:
        doc.add_heading('🔴 1. SATIŞ ÖNERİLERİ', level=1)
        
        total_sell_value = sum([r['recommended_value'] for r in sell_actions])
        doc.add_paragraph(f'Toplam Satış Değeri: {total_sell_value:,.2f} TL')
        doc.add_paragraph()
        
        for idx, rec in enumerate(sell_actions, 1):
            doc.add_heading(f'{idx}. {rec["symbol"]}', level=2)
            
            # İşlem bilgileri
            doc.add_paragraph(f'Önerilen İşlem: {rec["action"]}')
            doc.add_paragraph(f'Miktar: {rec["recommended_quantity"]:.0f} adet')
            doc.add_paragraph(f'İşlem Tutarı: {rec["recommended_value"]:,.2f} TL')
            doc.add_paragraph(f'Güncel Fiyat: {rec["current_price"]:,.2f} TL')
            
            if rec.get('avg_cost'):
                profit_loss = rec.get('profit_loss', 0)
                profit_loss_pct = rec.get('profit_loss_pct', 0)
                doc.add_paragraph(f'Kar/Zarar: {profit_loss:+,.2f} TL ({profit_loss_pct:+.1f}%)')
            
            if rec.get('action_reason'):
                doc.add_paragraph(f'💡 Sebep: {rec["action_reason"]}')
            
            doc.add_paragraph()
    
    # 2. YENİ ALIM ÖNERİLERİ
    if buy_actions:
        doc.add_heading('🟢 2. YENİ ALIM ÖNERİLERİ', level=1)
        
        total_buy_value = sum([r['recommended_value'] for r in buy_actions])
        doc.add_paragraph(f'Toplam Alım Değeri: {total_buy_value:,.2f} TL')
        doc.add_paragraph()
        
        for idx, rec in enumerate(buy_actions, 1):
            doc.add_heading(f'{idx}. {rec["symbol"]}', level=2)
            
            # İşlem bilgileri
            doc.add_paragraph(f'Önerilen Miktar: {rec["recommended_quantity"]:.0f} adet')
            doc.add_paragraph(f'Güncel Fiyat: {rec["current_price"]:,.2f} TL')
            doc.add_paragraph(f'İşlem Tutarı: {rec["recommended_value"]:,.2f} TL')
            
            if rec.get('target_price'):
                potential_return = ((rec['target_price'] - rec['current_price']) / rec['current_price']) * 100
                doc.add_paragraph(f'Hedef Fiyat: {rec["target_price"]:,.2f} TL (Potansiyel Getiri: %{potential_return:+.1f})')
            
            if rec.get('confidence'):
                doc.add_paragraph(f'Güven Skoru: %{rec["confidence"]*100:.0f}')
            
            # Neden AL dediğinin özeti
            if rec.get('action_reason'):
                doc.add_paragraph(f'💡 Ana Sebep: {rec["action_reason"]}')
            
            # Detaylı sebepler
            doc.add_paragraph('📋 Neden AL Önerisi:')
            reasons = generate_buy_reasons(rec)
            for reason in reasons:
                doc.add_paragraph(f'  • {reason}', style='List Bullet')
            
            doc.add_paragraph()
    
    # 3. POZİSYON ARTIRIMLARI
    # Sadece nakit varsa ve işlem tutarı > 0 olanları göster
    increase_actions_filtered = [r for r in increase_actions if r.get('recommended_value', 0) > 0]
    if increase_actions_filtered:
        doc.add_heading('📈 3. POZİSYON ARTIRIMLARI', level=1)
        
        for idx, rec in enumerate(increase_actions_filtered, 1):
            doc.add_heading(f'{idx}. {rec["symbol"]}', level=2)
            
            doc.add_paragraph(f'Mevcut Pozisyon: {rec["quantity"]:.0f} adet')
            doc.add_paragraph(f'Ek Önerilen: {rec["recommended_quantity"]:.0f} adet')
            doc.add_paragraph(f'Toplam Olacak: {rec["quantity"] + rec["recommended_quantity"]:.0f} adet')
            doc.add_paragraph(f'İşlem Tutarı: {rec["recommended_value"]:,.2f} TL')
            
            if rec.get('action_reason'):
                doc.add_paragraph(f'💡 Sebep: {rec["action_reason"]}')
            
            # Neden ARTIR dediğinin özeti
            doc.add_paragraph('📋 Neden ARTIR Önerisi:')
            reasons = generate_increase_reasons(rec)
            for reason in reasons:
                doc.add_paragraph(f'  • {reason}', style='List Bullet')
            
            doc.add_paragraph()
    
    # 4. TAKİP EDİLECEK HİSSELER
    if hold_actions:
        doc.add_heading('🟡 4. TAKİP EDİLECEK HİSSELER', level=1)
        
        for idx, rec in enumerate(hold_actions, 1):
            doc.add_heading(f'{idx}. {rec["symbol"]}', level=2)
            
            doc.add_paragraph(f'Mevcut Pozisyon: {rec["quantity"]:.0f} adet')
            doc.add_paragraph(f'Güncel Fiyat: {rec.get("current_price", 0):,.2f} TL')
            
            if rec.get('avg_cost'):
                profit_loss = rec.get('profit_loss', 0)
                profit_loss_pct = rec.get('profit_loss_pct', 0)
                doc.add_paragraph(f'Kar/Zarar: {profit_loss:+,.2f} TL ({profit_loss_pct:+.1f}%)')
            
            if rec.get('action_reason'):
                doc.add_paragraph(f'💡 Durum: {rec["action_reason"]}')
            
            doc.add_paragraph()
    
    # İŞLEM ÖZETİ
    doc.add_heading('📊 İşlem Özeti', level=1)
    
    total_buy = sum([r['recommended_value'] for r in buy_actions + increase_actions_filtered])
    total_sell = sum([r['recommended_value'] for r in sell_actions])
    net_cash = total_cash + total_sell - total_buy
    
    summary_table = doc.add_table(rows=5, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ('📊 Portföy Değeri', f'{portfolio_value:,.2f} TL'),
        ('💵 Başlangıç Nakit', f'{total_cash:,.2f} TL'),
        ('💰 Toplam Satış', f'+{total_sell:,.2f} TL'),
        ('💸 Toplam Alım', f'-{total_buy:,.2f} TL'),
        ('💵 Kalan Nakit', f'{net_cash:,.2f} TL')
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # Footer
    doc.add_paragraph('_' * 50)
    doc.add_paragraph(f'Rapor Oluşturulma Tarihi: {export_date.strftime("%d.%m.%Y %H:%M")}')
    doc.add_paragraph('🤖 Robot Portföy Yöneticisi - AI Destekli Yatırım Önerileri')
    doc.add_paragraph()
    doc.add_paragraph('⚠️ UYARI: Bu öneriler yatırım tavsiyesi niteliği taşımaz. Yatırım kararlarınızı kendi araştırmanızı yaparak alın.')
    
    return doc


def generate_buy_reasons(rec):
    """AL önerisi için sebepler listesi oluşturur"""
    reasons = []
    
    # Model tahmini
    confidence = rec.get('confidence', 0)
    prediction = rec.get('prediction')
    
    if prediction == 1 and confidence > 0.6:
        reasons.append(f'Güçlü AI model sinyali (%{confidence*100:.0f} güven)')
    elif prediction == 1 and confidence > 0.5:
        reasons.append(f'Orta seviye AI model sinyali (%{confidence*100:.0f} güven)')
    
    # Teknik analiz sinyalleri
    result = rec.get('result', {})
    if result:
        rsi = result.get('rsi', 50)
        if rsi < 30:
            reasons.append(f'RSI aşırı satım seviyesinde ({rsi:.1f}) - Yükseliş fırsatı')
        elif rsi < 40:
            reasons.append(f'RSI düşük seviyede ({rsi:.1f}) - Potansiyel yükseliş')
        
        trend = result.get('trend_strength', '')
        if trend == 'Yükseliş':
            reasons.append('Fiyat hareketli ortalamaların üzerinde - Güçlü trend')
        
        volume_ratio = result.get('volume_ratio', 1.0)
        if volume_ratio > 1.5:
            reasons.append(f'Hacim ortalamanın {volume_ratio:.1f}x üzerinde - İlgi artışı')
        
        price_change_1w = result.get('price_change_1w', 0)
        if price_change_1w > 3:
            reasons.append(f'Son 1 haftada %{price_change_1w:.1f} yükseliş - Momentum devam ediyor')
    
    # Hedef fiyat
    target_price = rec.get('target_price')
    current_price = rec.get('current_price', 0)
    if target_price and current_price:
        potential_return = ((target_price - current_price) / current_price) * 100
        if potential_return > 10:
            reasons.append(f'Potansiyel getiri %{potential_return:.1f} - Yüksek kazanç fırsatı')
        elif potential_return > 5:
            reasons.append(f'Potansiyel getiri %{potential_return:.1f} - İyi kazanç fırsatı')
    
    # Portföy tahsisi
    allocation = rec.get('allocation_pct', 0)
    if allocation > 0:
        reasons.append(f'Portföy tahsisi %{allocation*100:.0f} - Risk dağılımına uygun')
    
    # Eğer hiç sebep yoksa genel mesaj
    if not reasons:
        reasons.append('AI modeli pozitif sinyal veriyor')
    
    return reasons[:5]  # En fazla 5 sebep


def generate_increase_reasons(rec):
    """ARTIR önerisi için sebepler listesi oluşturur"""
    reasons = []
    
    confidence = rec.get('confidence', 0)
    if confidence > 0.6:
        reasons.append(f'Güçlü yükseliş sinyali (%{confidence*100:.0f} güven)')
    
    result = rec.get('result', {})
    if result:
        rsi = result.get('rsi', 50)
        if rsi < 30:
            reasons.append('Aşırı satım bölgesinde - İyi fırsat')
        
        trend = result.get('trend_strength', '')
        if trend == 'Yükseliş':
            reasons.append('Mevcut trend güçlü - Artırım mantıklı')
    
    profit_loss_pct = rec.get('profit_loss_pct', 0)
    if profit_loss_pct > -3:
        reasons.append('Mevcut pozisyon zararda değil - Artırım güvenli')
    
    return reasons[:4]  # En fazla 4 sebep

